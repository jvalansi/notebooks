import docx
import collections
from collections import defaultdict, Counter
import pprint
import os
import string
import re
import sys
import json
import pandas as pd
import numpy as np
import csv
from tqdm import tqdm_notebook as tqdm
from sklearn.decomposition import PCA
from sklearn.manifold import TSNE
from sklearn.cluster import KMeans
from sklearn.neighbors import NearestNeighbors, DistanceMetric
import matplotlib.pyplot as plt
from wordfreq import word_frequency, tokenize
from yandex_translate import YandexTranslate, YandexTranslateException
# import textblob
MIN_FREQ = 1e-08


def setText(fname, texts):
    doc = docx.Document()
    for text in texts:
        doc.add_paragraph(text)
    try:
        doc.save(fname)
    except (PermissionError, FileNotFoundError):
        print(fname)


def setTexts(path_name, text_dict):
    if not os.path.exists(path_name):
        os.makedirs(path_name)
    for fname,text in tqdm(text_dict.items()):
        fpath = os.path.join(path_name, fname+'.docx')
        setText(fpath, text)
        

def getText(filename):
    doc = docx.Document(filename)
    sentences = []
    for sentence in doc.paragraphs:
        sentences.append(sentence.text)#.strip())
    return sentences


def get_texts(path_name):
    indices=[]
    texts=[]
    for fname in os.listdir(path_name):
        fpath = os.path.join(path_name, fname)
        try:
            text = list(filter(None, getText(fpath)))
            indices.extend([fname]*len(text))
            texts.extend(text)
        except (docx.opc.exceptions.PackageNotFoundError, ValueError):
            print("{} Not added".format(fpath))
    return pd.DataFrame(texts, index=indices)


def translate_texts(texts):
    translate = YandexTranslate('trnsl.1.1.20181224T230635Z.693eab4c8554b089.84d0171c058028fc88a5905d980a60d305fdfad1')
    en_sents = []
    N = 10
    for i in tqdm(range(len(texts)//N+1)):
        try:
            en_sents.append(translate.translate(texts[i*N:(i+1)*N], 'en')['text'])
        except YandexTranslateException as e:
            print(i)
            print(e)    
    en_sents = [en_sent for sent in en_sents for en_sent in sent]
    return en_sents

    #en_blobs = []
    #for text in tqdm(texts):
    #    tb = textblob.TextBlob(text)
    #    try:
    #        en_blobs.append(tb.translate(to='en'))
    #    except textblob.exceptions.NotTranslated:
    #        en_blobs.append(text)
    #en_texts = [en_blob if type(en_blob) is str else en_blob.raw for en_blob in en_blobs]
    #return en_texts
            
def get_freqs(fpath='/Users/jordanvalansi/Downloads/he-2012/he.txt'):
    with open(fpath) as f:
        lines = f.readlines()
    freqs = {l.split()[0]: int(l.split()[1]) for l in lines}
    return freqs
    #for f in freqs:
    #    print(f)

def tf_idf(c, freqs):   
    d = {}
    for k, v in c.items():
        k = str(k)
        if k not in freqs:
            print("{} not in freqs".format(k))
            continue
        f = freqs[str(k)] if k in freqs else 1.0
        d[k] = float(v)/float(f)
    return collections.Counter(d)

def main():  
    path_name = '/Users/jordanvalansi/Downloads/Thoughts' 
    texts = get_texts(path_name)
    texts_words = [ re.split('[\s'+string.punctuation+']+', text) for text in texts]
    words = [w for text in texts_words for w in text]
    c = collections.Counter(words)
    c2 = {w: sum([w in text_words for text_words in texts_words]) for w in c}
    c2 = {k: v for k, v in c2.items() if v>3}
    #bigrams = [b for l in texts for b in zip(l.split(" ")[:-1], l.split(" ")[1:])]
    freqs = get_freqs()
    c3 = tf_idf(c2, freqs)
    for k,v in c3.most_common(100):
        print("{} {}".format(k,v))


def tokenize_he(word):
    """
    >>> tokenize('וכשמהבית')
    ['ו', 'כ', 'ש', 'מ', 'ה', 'בית']
    >>> tokenize('מה')
    ['מ', 'ה']
    >>> tokenize('המטרות')
    ['ה', 'מטרות']


    """
    tokens = []
    if not word: return tokens
    for letter in 'וכשמלבה':
        if word[0]!=letter:
            continue
        tokens.append(letter)
        word = word[1:]
    if word: tokens.append(word)
    return tokens
#     return list(filter(None, re.split('(^ו?כ?ש?מ?ל?ב?ה?)', word)))


def remove_prefix(word):
    """
    >>> remove_prefix('וכשמהבית')
    'בית'
    >>> remove_prefix('מה')
    ''
    """

    return re.sub('(^ו?כ?ש?מ?ל?ב?ה?)', '', word)

# import doctest
# doctest.testmod()
# import timeit
# timeit.timeit('[remove_prefix(word) for word in [\'וכשמהבית\', \'מה\', \'שלום\', \'חיים\']]' , setup="from __main__ import remove_prefix")


def get_top_words(words, lang='en', top=3):
    tf_idfs = {}
    c = Counter(words)
    for word, count in c.items():
        freq = word_frequency(word, lang)
        if not freq:
            continue
        tf_idfs[word] = count/freq
    if not tf_idfs:
        return [None]
    return sorted(tf_idfs.keys(), key=lambda k: tf_idfs[k], reverse=True)[:top]


def calc_pair_similarity(item0, item1):
    # intersection / union : if equal is 1 if no intersection 0
    set0 = set(item0)
    set1 = set(item1)
    intersection = set0.intersection(set1)
    union = set0.union(set1)
    return len(intersection)/len(union)


def words2one_hot(words, freqs):
    words = [choose_word(word, freqs) for word in words]
    return np.array([words.count(word)/freqs[word] for word in freqs])


def choose_word(word, lang='en'):
    lemma = remove_prefix(word)
    word_freq = word_frequency(word, lang)
    lemma_freq = word_frequency(lemma, lang)
    return lemma if lemma_freq>word_freq else word


# get tf_idf word in cluster
def gen_text_dict(clusters, texts_words, texts, lang='en'):
    text_dict = defaultdict(list)
    for name, cluster in clusters.items():
        words = [choose_word(word, lang) for i in cluster for word in texts_words[i]]
        uniqs = list(set(words).intersection(freqs))
        if not uniqs:
            continue
        tfs = [words.count(word)/np.log(word_frequency(word, lang)) for word in uniqs]
        text_name = uniqs[np.argmax(tfs)]
        text_dict[text_name].extend([texts[i] for i in cluster])
    return text_dict


def get_nearest_neighbour(distmat, i):
    return distmat[i].argsort()[:2][1]


def get_nn_dict(distmat):
    return dict(zip(range(len(distmat)), [get_nearest_neighbour(distmat, i) for i in range(len(distmat))]))

def get_nns(train, query=None, metric='cosine'):
    if query is None:
        query = train
    nn = NearestNeighbors(metric=metric).fit(train)
    _, indices = nn.kneighbors(query, n_neighbors=2)
    return [train.index.values[i][-1] for i in indices]

def get_cluster(d):
    item = next(iter(d))
    cluster = []
    while item in d:
        cluster.append(item)
        new_item = d[item]
        d.pop(item, None)
        item = new_item
    return cluster    


def get_vecs(fpath, sep=' '):
    d = {}
    with open(fpath) as f:
        lines = f.readlines()
        for line in tqdm(lines):
            line = line.strip()
            line = line.split(sep)
#             try:
            d[line[0]] = np.array(line[1:], dtype=float)
#             except ValueError:
#                 print(line)
#                 continue
    return d


def text2vec(text, w2v, lang='en'):
    vecs = []
    for word in text:
        f = word_frequency(word, lang)
        if not f or word.lower() not in w2v:
            continue
        vecs.append(w2v[word.lower()]/f)
    vecs_array = np.array(vecs)
    
    vec = vecs_array.mean(axis=0) if len(vecs_array.shape)>1 else w2v['.'] if vecs_array.shape[0]==0 else vecs_array#pd.Series(pd.concat([df.min(), df.mean(), df.max()], axis=0)).values
    return vec


def split_largest(orig_labels, orig_cluster):
    # find biggest cluster label (maybe should be largest number)
    max_label = orig_labels.max()
    # split biggest cluster
    kmeans = KMeans(n_clusters=2, random_state=0).fit(orig_cluster[orig_labels==max_label])
    # label cluster according to split
    labels_ = pd.Series(kmeans.labels_)
        # if 0 cluster is bigger than than 1, swap
    if len(labels_[labels_==0])>len(labels_[labels_==1]):
        labels_[labels_==0] = -1
        labels_[labels_==1] = 0
        labels_[labels_==-1] = 1
#         labels_ = np.abs(labels_ - 1)
    # get largest split, label as max+1
    orig_labels[orig_labels==max_label] = labels_.values + max_label
    return orig_labels


def get_tfidf(words):
    c = collections.Counter(words)
    count_df = pd.DataFrame(list(c.values()), index=list(c.keys()), columns=['count'])
    count_df['freq'] = count_df.apply(lambda x: word_frequency(x.name, 'en'), axis=1)
    count_df = count_df[count_df['freq']>0]
    count_df['tfidf'] = count_df.apply(lambda x: x['count']/x['freq'], axis=1)
    count_df= count_df.sort('tfidf', ascending=False)
    return count_df